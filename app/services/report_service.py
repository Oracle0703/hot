from __future__ import annotations

import hashlib
from datetime import datetime
from pathlib import Path
from urllib.parse import urlsplit
from uuid import UUID

import httpx
import portalocker
from docx import Document
from docx.shared import Inches
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.db import get_reports_root
from app.models.item import CollectedItem
from app.models.job import CollectionJob
from app.models.job_log import JobLog
from app.models.report import Report
from app.models.source import Source
from app.services.content_pipeline_service import ContentPipelineService
from app.services.network_access_policy import build_httpx_request_kwargs
from app.services.published_at_display import format_published_at


class ReportService:
    def __init__(self, session: Session, reports_root: Path | None = None) -> None:
        self.session = session
        self.reports_root = reports_root or get_reports_root()
        self.last_content_item_ids: list[UUID] = []

    def generate_for_job(self, job: CollectionJob, source_runs: list[dict[str, object]]) -> Report:
        report_dir = self.reports_root / "global"
        report_dir.mkdir(parents=True, exist_ok=True)
        markdown_path = report_dir / "hot-report.md"
        docx_path = report_dir / "hot-report.docx"
        markdown_temp_path = self._build_temp_report_path(markdown_path, job, "md")
        docx_temp_path = self._build_temp_report_path(docx_path, job, "docx")
        backup_paths = self._build_report_backup_paths(markdown_path, docx_path)
        existing_report_files = self._build_report_existing_states(markdown_path, docx_path)
        activation_started = False

        try:
            pipeline_result = ContentPipelineService(self.session).ingest_run(job.id, source_runs)
            self._upsert_collected_items(job, source_runs)
            self.session.flush()
            self.last_content_item_ids = [item.id for item in pipeline_result.content_items if getattr(item, "id", None) is not None]
            logs, items, sources = self._load_report_context(job)
            media_assets = self._sync_item_media_assets(report_dir, items)
            self._prune_unused_media_assets(report_dir, media_assets)
            markdown_content = self._build_global_markdown(job, items, sources, logs, media_assets, report_dir)
            markdown_temp_path.write_text(markdown_content, encoding="utf-8")
            self._write_docx(docx_temp_path, job, items, sources, logs, media_assets)
            activation_started = True
            self._activate_prepared_report_files(
                markdown_path=markdown_path,
                markdown_temp_path=markdown_temp_path,
                docx_path=docx_path,
                docx_temp_path=docx_temp_path,
                backup_paths=backup_paths,
            )

            report = self._get_or_create_global_report(job, markdown_path, docx_path)
            self.session.commit()
            self.session.refresh(report)
            return report
        except Exception:
            self.last_content_item_ids = []
            self.session.rollback()
            if activation_started:
                self._restore_report_files(markdown_path, docx_path, backup_paths, existing_report_files)
            raise
        finally:
            self._cleanup_report_paths(markdown_temp_path, docx_temp_path, *backup_paths.values())

    def list_reports(self, limit: int = 20) -> list[Report]:
        statement = select(Report).order_by(Report.created_at.desc()).limit(limit)
        return list(self.session.scalars(statement).all())

    def get_report(self, report_id: str) -> Report | None:
        return self.session.get(Report, UUID(report_id))

    def read_markdown(self, report: Report) -> str:
        return Path(report.markdown_path).read_text(encoding="utf-8")

    def clear_collected_items(self) -> int:
        result = self.session.execute(delete(CollectedItem))
        self.session.commit()
        return int(result.rowcount or 0)

    def _upsert_collected_items(self, job: CollectionJob, source_runs: list[dict[str, object]]) -> None:
        seen_at = job.finished_at or datetime.utcnow()
        for run in source_runs:
            source_id = run.get("source_id")
            if source_id is None:
                continue
            for item in run.get("items", []) or []:
                normalized_hash = self._compute_normalized_hash(str(source_id), item)
                collected_item = self.session.scalar(
                    select(CollectedItem).where(CollectedItem.normalized_hash == normalized_hash)
                )
                published_at = self._parse_published_at(item.get("published_at"))
                published_at_text = self._string_or_none(item.get("published_at_text")) or self._string_or_none(item.get("published_at"))
                image_urls = self._normalize_image_urls(item.get("image_urls"))
                cover_image_url = self._string_or_none(item.get("cover_image_url"))
                like_count = self._int_or_none(item.get("like_count"))
                reply_count = self._int_or_none(item.get("reply_count"))
                view_count = self._int_or_none(item.get("view_count"))
                if collected_item is None:
                    collected_item = CollectedItem(
                        source_id=source_id,
                        job_id=job.id,
                        first_seen_job_id=job.id,
                        last_seen_job_id=job.id,
                        title=self._normalize_title(item.get("title")) or "未命名帖子",
                        url=str(item.get("url") or ""),
                        author=self._string_or_none(item.get("author")),
                        published_at=published_at,
                        published_at_text=published_at_text,
                        first_seen_at=seen_at,
                        last_seen_at=seen_at,
                        heat_score=self._string_or_none(item.get("heat_score")),
                        cover_image_url=cover_image_url,
                        like_count=like_count,
                        reply_count=reply_count,
                        view_count=view_count,
                        excerpt=self._string_or_none(item.get("excerpt")),
                        image_urls=image_urls,
                        normalized_hash=normalized_hash,
                    )
                    self.session.add(collected_item)
                    continue

                collected_item.source_id = source_id
                collected_item.job_id = job.id
                collected_item.last_seen_job_id = job.id
                collected_item.last_seen_at = seen_at
                collected_item.title = self._normalize_title(item.get("title")) or collected_item.title or "未命名帖子"
                collected_item.url = str(item.get("url") or collected_item.url or "")
                collected_item.author = self._string_or_none(item.get("author")) or collected_item.author
                collected_item.published_at = published_at or collected_item.published_at
                collected_item.published_at_text = published_at_text or collected_item.published_at_text
                collected_item.heat_score = self._string_or_none(item.get("heat_score")) or collected_item.heat_score
                collected_item.cover_image_url = cover_image_url or collected_item.cover_image_url
                collected_item.like_count = like_count if like_count is not None else collected_item.like_count
                collected_item.reply_count = reply_count if reply_count is not None else collected_item.reply_count
                collected_item.view_count = view_count if view_count is not None else collected_item.view_count
                collected_item.excerpt = self._string_or_none(item.get("excerpt")) or collected_item.excerpt
                if image_urls:
                    collected_item.image_urls = image_urls

    def _compute_normalized_hash(self, source_id: str, item: dict[str, object]) -> str:
        url = str(item.get("url") or "").strip()
        if url:
            raw_value = f"{source_id}|{url}"
        else:
            title = str(item.get("title") or "").strip()
            published_at = str(item.get("published_at") or "").strip()
            raw_value = f"{source_id}|{title}|{published_at}"
        return hashlib.sha256(raw_value.encode("utf-8")).hexdigest()

    def _get_or_create_global_report(self, job: CollectionJob, markdown_path: Path, docx_path: Path) -> Report:
        reports = list(self.session.scalars(select(Report).order_by(Report.created_at.asc())).all())
        if reports:
            report = reports[0]
            for extra_report in reports[1:]:
                self.session.delete(extra_report)
        else:
            report = Report(job_id=job.id, markdown_path=str(markdown_path), docx_path=str(docx_path))
            self.session.add(report)
            self.session.flush()

        report.job_id = job.id
        report.markdown_path = str(markdown_path)
        report.docx_path = str(docx_path)
        return report

    def _load_report_context(self, job: CollectionJob) -> tuple[list[JobLog], list[CollectedItem], dict[str, Source]]:
        logs = self._list_job_logs(job.id)
        items = list(
            self.session.scalars(
                select(CollectedItem).order_by(CollectedItem.last_seen_at.desc(), CollectedItem.first_seen_at.desc())
            ).all()
        )
        source_ids = sorted({item.source_id for item in items}, key=str)
        sources = {
            str(source.id): source
            for source in self.session.scalars(select(Source).where(Source.id.in_(source_ids))).all()
        } if source_ids else {}
        return logs, items, sources

    def _build_global_markdown(
        self,
        job: CollectionJob,
        items: list[CollectedItem],
        sources: dict[str, Source],
        logs: list[JobLog],
        media_assets: dict[str, list[Path]],
        report_dir: Path,
    ) -> str:
        failed_source_ids = self._extract_failed_source_ids(logs)
        lines = [
            "# 热点采集报告",
            "",
            "## 任务概览",
            f"- 任务ID: {job.id}",
            f"- 状态: {job.status}",
            f"- 总来源: {job.total_sources}",
            f"- 成功来源: {job.success_sources}",
            f"- 失败来源: {job.failed_sources}",
            "",
            "## 热点帖子",
        ]

        if not items:
            lines.append("- 当前无历史帖子")
        else:
            grouped_items: dict[str, list[CollectedItem]] = {}
            for item in items:
                grouped_items.setdefault(str(item.source_id), []).append(item)

            for source_id in sorted(grouped_items, key=lambda current_id: sources.get(current_id).name if sources.get(current_id) else str(current_id)):
                source = sources.get(source_id)
                lines.append("")
                lines.append(f"### 来源: {source.name if source is not None else source_id}")
                source_items = sorted(
                    grouped_items[source_id],
                    key=lambda current_item: (
                        0 if current_item.first_seen_job_id == job.id else 1,
                        -(current_item.last_seen_at or datetime.min).timestamp(),
                    ),
                )
                for item in source_items:
                    published_at = self._format_datetime(item.published_at, getattr(item, "published_at_text", None))
                    prefix = ""
                    if item.first_seen_job_id == job.id:
                        prefix = "[NEW] "
                    elif item.last_seen_job_id != job.id and str(item.source_id) not in failed_source_ids:
                        prefix = "[本次未抓到] "
                    if item.url:
                        lines.append(f"- {prefix}[{self._normalize_title(item.title) or '未命名帖子'}]({item.url}) - {published_at}")
                    else:
                        lines.append(f"- {prefix}{self._normalize_title(item.title) or '未命名帖子'} - {published_at}")
                    for asset_path in media_assets.get(str(item.id), []):
                        lines.append(f"  ![]({asset_path.relative_to(report_dir).as_posix()})")

        lines.append("")
        lines.append("## 异常摘要")
        if not logs:
            lines.append("- 无异常")
        else:
            for log in logs:
                lines.append(f"- [{log.level}] {log.message}")

        lines.append("")
        return "\n".join(lines)

    def _list_job_logs(self, job_id) -> list[JobLog]:
        statement = select(JobLog).where(JobLog.job_id == job_id).order_by(JobLog.created_at.asc())
        return list(self.session.scalars(statement).all())

    def _parse_published_at(self, value: object) -> datetime | None:
        if isinstance(value, datetime):
            return value
        if value is None:
            return None
        text = str(value).strip()
        if not text:
            return None
        normalized = text.replace("Z", "+00:00")
        for parser in (datetime.fromisoformat,):
            try:
                return parser(normalized)
            except ValueError:
                continue
        for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
            try:
                return datetime.strptime(text, fmt)
            except ValueError:
                continue
        return None

    def _format_datetime(self, value: datetime | None, raw_text: object = None) -> str:
        return format_published_at(value, raw_text)

    def _string_or_none(self, value: object) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        return text or None

    def _normalize_title(self, value: object) -> str | None:
        text = self._string_or_none(value)
        if text is None:
            return None
        return " ".join(text.split())

    def _normalize_image_urls(self, value: object) -> list[str]:
        if not isinstance(value, list):
            return []
        normalized: list[str] = []
        for item in value:
            text = str(item or "").strip()
            if text and text not in normalized:
                normalized.append(text)
        return normalized

    def _int_or_none(self, value: object) -> int | None:
        try:
            if value is None or str(value).strip() == "":
                return None
            return int(value)
        except (TypeError, ValueError):
            return None

    def _build_temp_report_path(self, target_path: Path, job: CollectionJob, label: str) -> Path:
        return target_path.with_name(f".{target_path.name}.{job.id}.{label}.tmp")

    def _build_report_backup_paths(self, markdown_path: Path, docx_path: Path) -> dict[Path, Path]:
        return {
            markdown_path: markdown_path.with_name(f"{markdown_path.name}.bak"),
            docx_path: docx_path.with_name(f"{docx_path.name}.bak"),
        }

    def _build_report_existing_states(self, markdown_path: Path, docx_path: Path) -> dict[Path, bool]:
        return {
            markdown_path: markdown_path.exists(),
            docx_path: docx_path.exists(),
        }

    def _activate_prepared_report_files(
        self,
        markdown_path: Path,
        markdown_temp_path: Path,
        docx_path: Path,
        docx_temp_path: Path,
        backup_paths: dict[Path, Path],
    ) -> None:
        # TC-RPT-001: 跨进程/线程互斥，避免两个 writer 同时调度 .bak 与 .replace 导致中间态破坏。
        lock_path = markdown_path.with_name('.hot-report.lock')
        markdown_path.parent.mkdir(parents=True, exist_ok=True)
        with portalocker.Lock(str(lock_path), 'w', timeout=30):
            self._replace_report_file(markdown_path, markdown_temp_path, backup_paths[markdown_path])
            self._replace_report_file(docx_path, docx_temp_path, backup_paths[docx_path])

    def _replace_report_file(self, target_path: Path, temp_path: Path, backup_path: Path) -> None:
        backup_path.unlink(missing_ok=True)
        backup_created = False
        if target_path.exists():
            target_path.replace(backup_path)
            backup_created = True
        try:
            temp_path.replace(target_path)
        except Exception:
            if backup_created and backup_path.exists():
                target_path.unlink(missing_ok=True)
                backup_path.replace(target_path)
            raise

    def _restore_report_files(
        self,
        markdown_path: Path,
        docx_path: Path,
        backup_paths: dict[Path, Path],
        existing_report_files: dict[Path, bool],
    ) -> None:
        self._restore_report_file(markdown_path, backup_paths[markdown_path], existing_report_files[markdown_path])
        self._restore_report_file(docx_path, backup_paths[docx_path], existing_report_files[docx_path])

    def _restore_report_file(self, target_path: Path, backup_path: Path, existed_before_activation: bool) -> None:
        if backup_path.exists():
            target_path.unlink(missing_ok=True)
            backup_path.replace(target_path)
            return
        if not existed_before_activation:
            target_path.unlink(missing_ok=True)

    def _cleanup_report_paths(self, *paths: Path) -> None:
        for path in paths:
            path.unlink(missing_ok=True)

    def _sync_item_media_assets(self, report_dir: Path, items: list[CollectedItem]) -> dict[str, list[Path]]:
        assets_dir = report_dir / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)
        media_assets: dict[str, list[Path]] = {}

        for item in items:
            image_urls = list(getattr(item, "image_urls", []) or [])
            if not image_urls:
                continue
            saved_paths: list[Path] = []
            for index, image_url in enumerate(image_urls):
                destination_path = assets_dir / self._build_image_asset_name(image_url, index)
                try:
                    saved_path = self._download_image_asset(image_url, destination_path)
                except Exception:
                    continue
                saved_paths.append(saved_path)
            if saved_paths:
                media_assets[str(item.id)] = saved_paths

        return media_assets

    def _prune_unused_media_assets(self, report_dir: Path, media_assets: dict[str, list[Path]]) -> None:
        assets_dir = report_dir / "assets"
        if not assets_dir.exists():
            return

        active_assets = {path.resolve() for paths in media_assets.values() for path in paths}
        for asset_path in assets_dir.iterdir():
            if not asset_path.is_file():
                continue
            if asset_path.resolve() not in active_assets:
                asset_path.unlink(missing_ok=True)

    def _build_image_asset_name(self, image_url: str, index: int) -> str:
        suffix = Path(urlsplit(image_url).path).suffix or ".jpg"
        digest = hashlib.sha256(image_url.encode("utf-8")).hexdigest()[:16]
        return f"{digest}-{index}{suffix}"

    def _download_image_asset(self, image_url: str, destination_path: Path) -> Path:
        if destination_path.exists():
            return destination_path
        with httpx.Client(timeout=20.0, follow_redirects=True, **build_httpx_request_kwargs(image_url)) as client:
            response = client.get(image_url)
            response.raise_for_status()
        destination_path.parent.mkdir(parents=True, exist_ok=True)
        destination_path.write_bytes(response.content)
        return destination_path

    def _write_docx(
        self,
        output_path: Path,
        job: CollectionJob,
        items: list[CollectedItem],
        sources: dict[str, Source],
        logs: list[JobLog],
        media_assets: dict[str, list[Path]],
    ) -> None:
        failed_source_ids = self._extract_failed_source_ids(logs)
        document = Document()
        document.add_heading("热点采集报告", level=0)

        document.add_heading("任务概览", level=1)
        for line in (
            f"任务ID: {job.id}",
            f"状态: {job.status}",
            f"总来源: {job.total_sources}",
            f"成功来源: {job.success_sources}",
            f"失败来源: {job.failed_sources}",
        ):
            document.add_paragraph(line, style="List Bullet")

        document.add_heading("热点帖子", level=1)
        if not items:
            document.add_paragraph("当前无历史帖子")
        else:
            grouped_items: dict[str, list[CollectedItem]] = {}
            for item in items:
                grouped_items.setdefault(str(item.source_id), []).append(item)

            for source_id in sorted(grouped_items, key=lambda current_id: sources.get(current_id).name if sources.get(current_id) else str(current_id)):
                source = sources.get(source_id)
                document.add_heading(f"来源: {source.name if source is not None else source_id}", level=2)
                source_items = sorted(
                    grouped_items[source_id],
                    key=lambda current_item: (
                        0 if current_item.first_seen_job_id == job.id else 1,
                        -(current_item.last_seen_at or datetime.min).timestamp(),
                    ),
                )
                for item in source_items:
                    prefix = ""
                    if item.first_seen_job_id == job.id:
                        prefix = "[NEW] "
                    elif item.last_seen_job_id != job.id and str(item.source_id) not in failed_source_ids:
                        prefix = "[本次未抓到] "
                    published_at = self._format_datetime(item.published_at, getattr(item, "published_at_text", None))
                    document.add_paragraph(f"{prefix}{self._normalize_title(item.title) or '未命名帖子'} - {published_at}", style="List Bullet")
                    if item.url:
                        document.add_paragraph(item.url)
                    for asset_path in media_assets.get(str(item.id), []):
                        try:
                            document.add_picture(str(asset_path), width=Inches(5.5))
                        except Exception:
                            continue

        document.add_heading("异常摘要", level=1)
        if not logs:
            document.add_paragraph("无异常", style="List Bullet")
        else:
            for log in logs:
                document.add_paragraph(f"[{log.level}] {log.message}", style="List Bullet")

        document.save(output_path)

    def _extract_failed_source_ids(self, logs: list[JobLog]) -> set[str]:
        return {
            str(log.source_id)
            for log in logs
            if log.level == "error" and log.source_id is not None
        }
