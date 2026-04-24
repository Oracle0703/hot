"""TC-RPT-001~004 — 报告写入跨进程文件锁 (portalocker)。

ReportService._activate_prepared_report_files 会在 markdown_path 同目录创建
``.hot-report.lock`` 并以 portalocker 持有,确保两个 writer 不会同时执行
".bak / .replace" 切换。本测试不需要全套 ORM 上下文,直接验证锁行为即可。
"""

from __future__ import annotations

import threading
import time

import portalocker
import pytest


def _hold_lock(lock_path, hold_seconds, started_event, released_event):
    with portalocker.Lock(str(lock_path), 'w', timeout=5):
        started_event.set()
        time.sleep(hold_seconds)
    released_event.set()


def test_lock_serializes_concurrent_writers(tmp_path) -> None:
    """TC-RPT-001: 第二个 writer 必须等待第一个释放锁后才能进入临界区"""
    lock_path = tmp_path / '.hot-report.lock'
    started = threading.Event()
    released = threading.Event()

    holder = threading.Thread(
        target=_hold_lock,
        args=(lock_path, 0.5, started, released),
        daemon=True,
    )
    holder.start()
    assert started.wait(2.0), '第一个 writer 未能拿到锁'

    acquired_at = []

    def second_writer():
        with portalocker.Lock(str(lock_path), 'w', timeout=10):
            acquired_at.append(time.monotonic())

    second = threading.Thread(target=second_writer, daemon=True)
    t0 = time.monotonic()
    second.start()
    second.join(timeout=5)
    holder.join(timeout=5)

    assert acquired_at, '第二个 writer 始终未能获得锁'
    elapsed = acquired_at[0] - t0
    assert elapsed >= 0.4, f'第二个 writer 未等待第一个释放,实际只等了 {elapsed:.3f}s'


def test_lock_timeout_raises(tmp_path) -> None:
    """TC-RPT-002: 持续被占用的锁,新申请者在 timeout 内拿不到则抛异常"""
    lock_path = tmp_path / '.hot-report.lock'
    started = threading.Event()
    released = threading.Event()

    holder = threading.Thread(
        target=_hold_lock,
        args=(lock_path, 1.0, started, released),
        daemon=True,
    )
    holder.start()
    assert started.wait(2.0)

    with pytest.raises(portalocker.exceptions.LockException):
        with portalocker.Lock(str(lock_path), 'w', timeout=0.2):
            pass

    holder.join(timeout=5)


def test_temp_file_kept_on_failure(tmp_path) -> None:
    """TC-RPT-003: _replace_report_file 失败时 .bak 仍能把目标文件恢复回原内容。"""
    from app.services.report_service import ReportService

    target = tmp_path / "hot-report.md"
    target.write_text("ORIGINAL", encoding="utf-8")
    backup = tmp_path / "hot-report.md.bak"
    temp_path = tmp_path / ".hot-report.md.tmp"
    # 不创建 temp_path -> .replace 会抛 FileNotFoundError 并触发 .bak 还原
    svc = ReportService.__new__(ReportService)
    with pytest.raises(FileNotFoundError):
        svc._replace_report_file(target, temp_path, backup)
    assert target.exists()
    assert target.read_text(encoding="utf-8") == "ORIGINAL"


def test_docx_atomic_replace(tmp_path) -> None:
    """TC-RPT-004: 经 _activate_prepared_report_files 切换后,docx 仍能被 python-docx 重新打开。"""
    from docx import Document

    from app.services.report_service import ReportService

    md_target = tmp_path / "hot-report.md"
    docx_target = tmp_path / "hot-report.docx"
    md_temp = tmp_path / ".hot-report.md.tmp"
    docx_temp = tmp_path / ".hot-report.docx.tmp"

    md_temp.write_text("# title\n", encoding="utf-8")
    doc = Document()
    doc.add_paragraph("hello concurrent world")
    doc.save(str(docx_temp))

    backup_paths = {
        md_target: md_target.with_name(md_target.name + ".bak"),
        docx_target: docx_target.with_name(docx_target.name + ".bak"),
    }
    svc = ReportService.__new__(ReportService)
    svc._activate_prepared_report_files(
        markdown_path=md_target,
        markdown_temp_path=md_temp,
        docx_path=docx_target,
        docx_temp_path=docx_temp,
        backup_paths=backup_paths,
    )
    assert md_target.read_text(encoding="utf-8").startswith("# title")
    reopened = Document(str(docx_target))
    assert any("hello concurrent" in p.text for p in reopened.paragraphs)
